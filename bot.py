import os
import re
import pdfplumber
from docx import Document
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, CallbackQueryHandler, ContextTypes

TOKEN = os.getenv("BOT_TOKEN")

DOWNLOADS_DIR = 'downloads'
RESULTS_DIR = 'results'
TEMPLATE_PATH = 'template.docx'

os.makedirs(DOWNLOADS_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)

user_data = {}

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Отправьте PDF программы проверок")

async def handle_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    document = update.message.document

    if document.mime_type != 'application/pdf':
        await update.message.reply_text("Нужен PDF")
        return

    file_path = os.path.join(DOWNLOADS_DIR, document.file_name)
    file = await document.get_file()
    await file.download_to_drive(file_path)

    await update.message.reply_text("Файл получен, анализирую...")

    data = extract_data_from_pdf(file_path)

    user_id = update.message.from_user.id
    user_data[user_id] = data

    await ask_signer(update)

def extract_data_from_pdf(file_path):
    full_text = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            full_text += text + "\n"

    case_id = re.search(r'\d{2}-\d{2}-\d{6,8}', full_text)
    object_name = re.search(r'«([^»]+)»', full_text)
    date = re.search(r'([0-3]?\d\.[0-1]?\d\.\d{4})', full_text)
    developer = re.search(r'Застройщик:\s*(.+)', full_text)

    return {
        'case_id': case_id.group() if case_id else '',
        'object_name': object_name.group(1) if object_name else '',
        'date': date.group() if date else '',
        'developer': developer.group(1).strip() if developer else ''
    }

async def ask_signer(update):
    keyboard = [
        [InlineKeyboardButton("Начальник отдела", callback_data='signer_Начальник отдела')],
        [InlineKeyboardButton("Заместитель", callback_data='signer_Заместитель')]
    ]
    await update.message.reply_text("Выберите подписанта:", reply_markup=InlineKeyboardMarkup(keyboard))

async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    signer = query.data.replace('signer_', '')

    user_data[user_id]['signer'] = signer

    await query.edit_message_text(f"Подписант: {signer}")
    await query.message.reply_text("Введите исполнителя:")

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id

    if user_id not in user_data:
        await update.message.reply_text("Сначала отправьте PDF")
        return

    user_data[user_id]['executor'] = update.message.text

    data = user_data[user_id]

    output_path = os.path.join(RESULTS_DIR, f"result_{user_id}.docx")

    create_docx(data, output_path)

    with open(output_path, "rb") as f:
        await update.message.reply_document(f)

    del user_data[user_id]

def create_docx(data, output_path):
    doc = Document(TEMPLATE_PATH)

    for p in doc.paragraphs:
        for key, val in data.items():
            p.text = p.text.replace(f"{{{{{key}}}}}", val)

    doc.add_paragraph(f"Подписант: {data.get('signer','')}")
    doc.add_paragraph(f"Исполнитель: {data.get('executor','')}")

    doc.save(output_path)

def main():
    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.PDF, handle_pdf))
    app.add_handler(CallbackQueryHandler(button_callback))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    app.run_polling()

if __name__ == "__main__":
    main()
